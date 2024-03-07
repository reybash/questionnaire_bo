import re

from docx import Document
from io import BytesIO


def input_file(file_name):
    """
       Reads a text file containing questions
       and durations, and returns a dictionary.

       The input file should follow the
       specified pattern, for example:
       1. How do you define system requirements based
       on the most successful ways in your experience?
          What methods and tools do you use for this? - 2 minutes
       ...

       Each question should be numbered and
       followed by its text and duration in minutes.

       """
    # Read the content of the file
    with open(file_name, "r", encoding="utf-8") as file:
        text = file.read()

    # Use regular expressions to extract questions and their duration
    pattern = re.compile(r'(\d+)\.(.*?)–?\s*(\d+)\s*(?:минуты|минут)',
                         re.DOTALL | re.UNICODE)
    matches = pattern.finditer(text)

    # Create a dictionary with questions and their duration
    questions_data = {}
    for match in matches:
        question_number, question_text, duration_minutes = match.groups()
        duration_seconds = int(duration_minutes) * 60
        # Convert minutes to seconds
        questions_data[int(question_number)] = {
            'text': f"{question_number}. {question_text.strip()} "
                    f"({duration_minutes} minute(s) to answer)",
            'duration': duration_seconds}

    return questions_data


def output_word_file(first_last_name, username, data_answers):
    word_file = BytesIO()
    doc = Document()

    # Add a heading for the report
    doc.add_heading(
        f"Report on answers from {first_last_name}", level=1)

    # Add user information to the document
    doc.add_paragraph(f"Telegram username: {username}")

    # Add a heading for answers
    doc.add_heading("Answers to questions", level=2)

    # Loop through questions and answers and add them to the document
    for question, answer in data_answers.items():
        p_question = doc.add_paragraph()
        run_question = p_question.add_run(f"Question: {question}")
        run_question.bold = True

        doc.add_paragraph(f"Answer: {answer}")
        doc.add_paragraph()  # Add an empty paragraph between questions

    doc.save(word_file)
    word_file.seek(0)

    return word_file.read()


def output_all_data_file(data):
    word_file = BytesIO()
    doc = Document()

    # Add a heading for the overall report
    doc.add_heading("Report on answers from all users",
                    level=1)

    current_user_id = None

    # Write data to the document
    for row in data:
        user_id = row[0]  # The first column corresponds to user_id

        # Check if the user has changed
        if user_id != current_user_id:
            # Add user information
            doc.add_heading(f"User ID: {user_id}", level=1)
            doc.add_paragraph(f"First and Last Name: {row[1]}")
            doc.add_paragraph(
                f"Telegram username: {row[2]}")
            doc.add_paragraph(f"Registration date: {row[3]}")

            # Update the current user
            current_user_id = user_id

        p_question = doc.add_paragraph()
        run_question = p_question.add_run(f"Question: {row[5]}")
        run_question.bold = True

        doc.add_paragraph(f"Answer: {row[6]}")
        doc.add_paragraph()  # Add an empty paragraph between entries

    doc.save(word_file)
    word_file.seek(0)

    return word_file.read()
